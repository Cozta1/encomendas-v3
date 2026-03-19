from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Estilos base ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0x76)
    hs.font.name = 'Calibri'


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    shading = run._element.get_or_add_rPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F0F0'
    })
    shading.append(shd)


# ═══════════════════════════════════════════
# CABEÇALHO
# ═══════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Relatório Técnico')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0x76)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Deploy do Sistema de Encomendas no Servidor')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

meta = doc.add_table(rows=3, cols=2)
meta.style = 'Table Grid'
meta_data = [
    ('Data', '12/03/2026'),
    ('Assunto', 'Solicitação de instalação do Docker Desktop e PostgreSQL no servidor da farmácia para deploy do Sistema de Encomendas'),
    ('Destinatário', 'João (Gestor de TI)')
]
for i, (k, v) in enumerate(meta_data):
    meta.rows[i].cells[0].text = k
    meta.rows[i].cells[1].text = v
    for p in meta.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
    for p in meta.rows[i].cells[1].paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)

doc.add_paragraph()

# ═══════════════════════════════════════════
# SEÇÃO 1 — Visão Geral
# ═══════════════════════════════════════════

doc.add_heading('1. Visão Geral da Arquitetura', level=1)

doc.add_paragraph(
    'A aplicação será implantada usando uma arquitetura híbrida, que combina o melhor de dois mundos:'
)

bullets = [
    ('Banco de dados (PostgreSQL 16)', 'instalado nativamente no Windows Server como um serviço do sistema. Isso garante que os dados fiquem armazenados diretamente no disco do servidor, com acesso direto, backups simples e total independência do Docker.'),
    ('Aplicação (Backend + Frontend)', 'executada dentro de containers Docker. Isso evita a necessidade de instalar Java, Node.js, Nginx e suas dependências diretamente no servidor, mantendo o ambiente limpo e isolado.'),
]
for title, desc in bullets:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Por que esta abordagem:', level=2)

reasons = [
    ('Segurança dos dados', 'Os dados do banco ficam em uma pasta comum do Windows (ex: D:\\PostgreSQL\\data), acessíveis diretamente pelo sistema de ficheiros. Não dependem de discos virtuais ou da VM do Docker.'),
    ('Backup simplificado', 'Backups podem ser feitos com ferramentas padrão do Windows e do PostgreSQL, sem necessidade de comandos Docker.'),
    ('Isolamento da aplicação', 'O backend (Java) e o frontend (Angular) ficam encapsulados em containers, sem instalar software adicional no sistema operacional.'),
    ('Resiliência', 'Se o Docker for atualizado, reinstalado ou tiver qualquer problema, os dados do banco de dados permanecem intactos, pois estão fora do Docker.'),
]
for title, desc in reasons:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(desc)

# ═══════════════════════════════════════════
# SEÇÃO 2 — Software Necessário
# ═══════════════════════════════════════════

doc.add_heading('2. Software Necessário para Instalação', level=1)

doc.add_paragraph('Serão necessários dois softwares instalados no servidor:')

doc.add_heading('2.1. PostgreSQL 16 (Banco de Dados)', level=2)

add_table(doc,
    ['Item', 'Detalhe'],
    [
        ['O que é', 'Sistema de banco de dados relacional, open-source e gratuito'],
        ['Download', 'https://www.postgresql.org/download/windows/'],
        ['Versão', 'PostgreSQL 16 (última versão estável)'],
        ['Instalador', 'Wizard gráfico padrão do Windows (próximo, próximo, concluir)'],
        ['Espaço em disco', '~250 MB (instalação) + espaço dos dados'],
        ['RAM estimada', '~40–50 MB em repouso (medido em testes)'],
        ['Serviço Windows', 'Instala automaticamente como serviço postgresql-x64-16'],
        ['Porta', '5432 (padrão, apenas localhost)'],
    ]
)

p = doc.add_paragraph()
run = p.add_run('Durante a instalação será solicitado:')
run.bold = True

install_steps = [
    'Pasta de dados — recomenda-se utilizar uma pasta fora do disco do sistema, como D:\\PostgreSQL\\data, para facilitar backups e evitar perda de dados em caso de reinstalação do Windows.',
    'Senha do utilizador postgres — esta senha será configurada no ficheiro .env da aplicação.',
    'Porta — manter a padrão (5432).',
    'Locale — manter o padrão ou selecionar Portuguese, Brazil.',
]
for b in install_steps:
    doc.add_paragraph(b, style='List Number')

doc.add_paragraph('Após a instalação, será necessário criar o banco de dados da aplicação:')
add_code_block(doc, 'psql -U postgres\nCREATE DATABASE sistema_encomendas;\n\\q')

doc.add_heading('2.2. Docker Desktop (Aplicação)', level=2)

add_table(doc,
    ['Item', 'Detalhe'],
    [
        ['O que é', 'Plataforma para executar aplicações em containers isolados'],
        ['Download', 'https://www.docker.com/products/docker-desktop/'],
        ['Versão', 'Docker Desktop (última versão estável)'],
        ['Instalador', 'Wizard gráfico padrão do Windows'],
        ['Espaço em disco', '~2–3 GB (instalação) + ~525 MB (imagens da aplicação)'],
        ['RAM estimada', '~620 MB em uso (medido em testes)'],
        ['Pré-requisito', 'WSL2 (o instalador do Docker configura automaticamente)'],
    ]
)

# ═══════════════════════════════════════════
# SEÇÃO 3 — Containers
# ═══════════════════════════════════════════

doc.add_heading('3. Containers Docker — Quantidade e Função', level=1)

doc.add_paragraph('Com o banco de dados fora do Docker, a aplicação utiliza apenas 2 (dois) containers:')

add_table(doc,
    ['#', 'Container', 'Imagem Base', 'Função'],
    [
        ['1', 'backend', 'eclipse-temurin:21-jre-alpine', 'API REST (Spring Boot / Java 21) — lógica de negócio, autenticação JWT, WebSocket em tempo real, envio de e-mails. Executa como utilizador não-root por segurança.'],
        ['2', 'frontend', 'nginx:alpine', 'Interface web (Angular + Nginx) — serve os ficheiros da aplicação web e atua como reverse proxy. Único ponto de acesso pela rede.'],
    ]
)

doc.add_heading('Diagrama de comunicação', level=2)

diagram = """  Navegador do utilizador
          |
          v (porta 8090)
    +--------------+
    |  Frontend    |  (Nginx - interface web)          -- DOCKER
    |  (Angular)   |
    +------+-------+
           | proxy interno (80 -> 8080)
           v
    +--------------+
    |  Backend     |  (Spring Boot - API e WebSocket)  -- DOCKER
    |  (Java 21)   |
    +------+-------+
           | conexao TCP (porta 5432)
           v
    +--------------+
    | PostgreSQL   |  (Banco de dados nativo)          -- WINDOWS
    |    16        |
    +--------------+"""

add_code_block(doc, diagram)

p = doc.add_paragraph()
run = p.add_run('Nota: ')
run.bold = True
p.add_run('O backend conecta ao PostgreSQL nativo do servidor via host.docker.internal:5432 — um endereço especial do Docker que aponta para o sistema operacional do servidor.')

# ═══════════════════════════════════════════
# SEÇÃO 4 — Recursos
# ═══════════════════════════════════════════

doc.add_heading('4. Consumo Real de Recursos (Medido em Testes)', level=1)

p = doc.add_paragraph()
run = p.add_run('Nota: ')
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0x76)
p.add_run('Os valores abaixo foram obtidos através de testes reais executados em 12/03/2026, sob carga simulada (múltiplas requisições simultâneas de autenticação e navegação).')

doc.add_heading('Memória RAM', level=2)

add_table(doc,
    ['Componente', 'Em Repouso (medido)', 'Sob Carga (medido)'],
    [
        ['PostgreSQL (nativo)', '~42 MB', '~43 MB'],
        ['Backend (Docker)', '604 MB', '609 MB'],
        ['Frontend (Docker)', '6 MB', '7 MB'],
        ['TOTAL', '~652 MB', '~659 MB'],
    ]
)

p = doc.add_paragraph()
run = p.add_run('Resumo: ')
run.bold = True
p.add_run('A aplicação consome aproximadamente 660 MB de RAM no total. O backend Java é o maior consumidor (~610 MB), o que é normal para aplicações Spring Boot — a memória é alocada no arranque e permanece estável. Recomenda-se que o servidor tenha pelo menos 1 GB de RAM livre.')

doc.add_heading('CPU', level=2)

add_table(doc,
    ['Componente', 'Em Repouso (medido)', 'Sob Carga (medido)'],
    [
        ['PostgreSQL (nativo)', '0,00%', '0,01%'],
        ['Backend (Docker)', '0,28%', '2,84%'],
        ['Frontend (Docker)', '0,00%', '0,00%'],
    ]
)

doc.add_paragraph('O impacto no processador é insignificante. Mesmo sob carga, o consumo máximo foi de apenas 2,84%.')

doc.add_heading('Espaço em Disco', level=2)

add_table(doc,
    ['Componente', 'Tamanho'],
    [
        ['PostgreSQL 16 (instalação)', '~250 MB'],
        ['Docker Desktop (instalação)', '~2–3 GB'],
        ['Imagem Backend (Spring Boot)', '428 MB'],
        ['Imagem Frontend (Angular + Nginx)', '96 MB'],
        ['Banco de dados (dados atuais)', '8,7 MB (crescimento gradual)'],
        ['TOTAL INICIAL', '~3,8 GB'],
    ]
)

p = doc.add_paragraph()
run = p.add_run('Nota: ')
run.bold = True
p.add_run('Sem o container do PostgreSQL no Docker, economiza-se 395 MB de espaço em imagens.')

# ═══════════════════════════════════════════
# SEÇÃO 5 — Portas
# ═══════════════════════════════════════════

doc.add_heading('5. Portas de Rede Utilizadas', level=1)

add_table(doc,
    ['Porta', 'Protocolo', 'Acessível externamente?', 'Utilização'],
    [
        ['8090', 'HTTP/TCP', 'Sim (única porta exposta)', 'Acesso à aplicação web pelo navegador'],
        ['8080', 'HTTP/TCP', 'Não (rede interna Docker)', 'API backend (Spring Boot)'],
        ['5432', 'TCP', 'Não (apenas localhost)', 'PostgreSQL nativo (banco de dados)'],
    ]
)

port_details = [
    'Porta 8090: Única porta acessível pela rede. Configurável via ficheiro .env (variável APP_PORT).',
    'Porta 8080: Interna ao Docker, invisível para o servidor e rede local.',
    'Porta 5432: O PostgreSQL nativo escuta apenas em localhost por padrão — não é acessível pela rede local nem pela internet, apenas pelo próprio servidor.',
    'Nenhuma dessas portas conflita com as portas padrão do MySQL (3306) utilizado pelo sistema da farmácia.',
]
for b in port_details:
    doc.add_paragraph(b, style='List Bullet')

# ═══════════════════════════════════════════
# SEÇÃO 6 — Armazenamento
# ═══════════════════════════════════════════

doc.add_heading('6. Armazenamento dos Dados', level=1)

doc.add_paragraph('Com o PostgreSQL instalado nativamente, os dados ficam em pastas comuns do Windows:')

add_table(doc,
    ['Dados', 'Localização no servidor'],
    [
        ['Banco de dados', 'Pasta configurada na instalação (ex: D:\\PostgreSQL\\data)'],
        ['Uploads da aplicação', 'Volume Docker: enc_uploads (dentro do Docker)'],
    ]
)

doc.add_heading('Vantagens do banco nativo:', level=2)

native_advantages = [
    'Os dados ficam diretamente no disco do servidor, sem camadas de virtualização.',
    'Acessíveis pelo Windows Explorer como qualquer outra pasta.',
    'Podem ser incluídos em rotinas de backup existentes do servidor (ex: backup de disco, cópia para rede).',
    'Independem totalmente do Docker — se o Docker for removido, atualizado ou falhar, os dados do banco permanecem intactos.',
    'Podem ser administrados diretamente via terminal com o comando psql (incluído na instalação do PostgreSQL).',
]
for b in native_advantages:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph('Exemplo de acesso ao banco via terminal:')
add_code_block(doc, 'psql -U postgres -d sistema_encomendas')

# ═══════════════════════════════════════════
# SEÇÃO 7 — Backup e Restauração
# ═══════════════════════════════════════════

doc.add_heading('7. Estratégia de Backup e Restauração', level=1)

doc.add_heading('7.1. Backup do Banco de Dados', level=2)

doc.add_heading('Backup manual (sob demanda):', level=3)
add_code_block(doc, 'pg_dump -U postgres -F c -f "D:\\Backups\\encomendas_backup.dump" sistema_encomendas')
doc.add_paragraph('Este comando gera um ficheiro .dump comprimido com todo o conteúdo do banco. O formato custom (-F c) é o mais eficiente e permite restauração seletiva de tabelas.')

doc.add_heading('Backup em formato SQL legível (alternativa):', level=3)
add_code_block(doc, 'pg_dump -U postgres sistema_encomendas > "D:\\Backups\\encomendas_backup.sql"')

doc.add_heading('Backup automatizado (recomendado):', level=3)
doc.add_paragraph('Criar um ficheiro backup_encomendas.bat na pasta de backups:')

bat_script = """@echo off
set PGPASSWORD=SUA_SENHA_AQUI
set BACKUP_DIR=D:\\Backups\\Encomendas
set DATA=%date:~6,4%%date:~3,2%%date:~0,2%

if not exist "%BACKUP_DIR%" mkdir "%BACKUP_DIR%"

pg_dump -U postgres -F c -f "%BACKUP_DIR%\\encomendas_%DATA%.dump" sistema_encomendas

:: Apagar backups com mais de 30 dias
forfiles /p "%BACKUP_DIR%" /s /m *.dump /d -30 /c "cmd /c del @path" 2>nul

echo Backup concluido: %BACKUP_DIR%\\encomendas_%DATA%.dump"""

add_code_block(doc, bat_script)

doc.add_paragraph('Depois, agendar no Task Scheduler do Windows:')
scheduler_steps = [
    'Abrir Task Scheduler → Create Basic Task',
    'Nome: "Backup Encomendas"',
    'Trigger: Diariamente, às 02:00 (ou horário de menor uso)',
    'Action: Executar o ficheiro backup_encomendas.bat',
]
for s in scheduler_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('Verificação de integridade do backup:', level=3)
add_code_block(doc, 'pg_restore -l "D:\\Backups\\Encomendas\\encomendas_20260312.dump"')
doc.add_paragraph('Este comando lista o conteúdo do backup sem restaurá-lo, permitindo verificar que o ficheiro está íntegro.')

doc.add_heading('7.2. Restauração do Banco de Dados', level=2)

doc.add_heading('Restauração completa (substitui todos os dados):', level=3)
add_code_block(doc, 'pg_restore -U postgres -d sistema_encomendas --clean --if-exists "D:\\Backups\\Encomendas\\encomendas_20260312.dump"')

doc.add_heading('Restauração a partir de ficheiro SQL:', level=3)
add_code_block(doc, 'psql -U postgres -d sistema_encomendas < "D:\\Backups\\encomendas_backup.sql"')

doc.add_heading('7.3. Backup dos Uploads', level=2)
doc.add_paragraph('Os ficheiros de upload ficam num volume Docker. Para copiar:')
add_code_block(doc, 'docker cp enc-backend-1:/app/uploads "D:\\Backups\\encomendas_uploads"')

# ═══════════════════════════════════════════
# SEÇÃO 8 — Inicialização Automática
# ═══════════════════════════════════════════

doc.add_heading('8. Inicialização Automática', level=1)

doc.add_paragraph('Ambos os componentes iniciam automaticamente com o servidor:')

add_table(doc,
    ['Componente', 'Mecanismo de auto-start'],
    [
        ['PostgreSQL', 'Serviço Windows postgresql-x64-16 — inicia automaticamente com o Windows (configurado na instalação)'],
        ['Backend + Frontend', 'Docker Desktop inicia com o Windows → containers com restart: unless-stopped sobem automaticamente'],
    ]
)

doc.add_heading('Ordem de arranque:', level=2)
boot_order = [
    'Windows inicia → PostgreSQL (serviço Windows) inicia automaticamente',
    'Docker Desktop inicia → container backend sobe e conecta ao PostgreSQL',
    'Backend fica saudável (health check) → container frontend sobe',
]
for b in boot_order:
    doc.add_paragraph(b, style='List Number')

doc.add_heading('Em caso de falha:', level=2)
failure_handling = [
    'Se o PostgreSQL parar, o serviço Windows pode ser configurado para reiniciar automaticamente (opção "Recovery" nas propriedades do serviço).',
    'Se um container Docker parar, a política restart: unless-stopped o reinicia automaticamente.',
]
for b in failure_handling:
    doc.add_paragraph(b, style='List Bullet')

# ═══════════════════════════════════════════
# SEÇÃO 9 — Independência
# ═══════════════════════════════════════════

doc.add_heading('9. Independência do Sistema da Farmácia', level=1)

p = doc.add_paragraph()
run = p.add_run('A aplicação de encomendas é 100% independente do sistema da farmácia.')
run.bold = True

add_table(doc,
    ['Aspeto', 'Sistema da Farmácia', 'Sistema de Encomendas'],
    [
        ['Banco de dados', 'MySQL (existente)', 'PostgreSQL 16 (novo, instalação separada)'],
        ['Porta do banco', '3306', '5432 (apenas localhost, sem conflito)'],
        ['Porta da aplicação', 'Portas próprias', '8090 (configurável)'],
        ['Armazenamento de dados', 'Disco do servidor', 'Pasta separada no disco (ex: D:\\PostgreSQL)'],
        ['Processos da aplicação', 'Serviços Windows', 'Containers Docker isolados'],
        ['Rede', 'Rede do servidor', 'Rede interna Docker + porta 8090'],
    ]
)

doc.add_heading('Garantias de não interferência:', level=2)
guarantees = [
    'Não há comunicação entre os dois sistemas. O sistema de encomendas não acessa, lê, modifica ou se conecta ao banco MySQL da farmácia em nenhuma circunstância.',
    'O PostgreSQL e o MySQL são bancos de dados completamente diferentes, com portas diferentes (5432 vs 3306), processos diferentes e dados em pastas diferentes.',
    'Não há instalação de Java, Node.js ou Nginx no sistema operacional — ficam apenas dentro dos containers Docker.',
    'Em caso de problema, basta parar os containers (docker compose down) e o serviço do PostgreSQL. O servidor volta a funcionar exactamente como antes. Não há efeitos colaterais.',
]
for b in guarantees:
    doc.add_paragraph(b, style='List Bullet')

# ═══════════════════════════════════════════
# SEÇÃO 10 — Resumo Executivo
# ═══════════════════════════════════════════

doc.add_heading('10. Resumo Executivo', level=1)

add_table(doc,
    ['Item', 'Resposta'],
    [
        ['Software a instalar', 'PostgreSQL 16 (banco de dados) + Docker Desktop (aplicação)'],
        ['Containers Docker', '2 (API backend + frontend web)'],
        ['RAM total medida', '~660 MB (recomendado 1 GB livre)'],
        ['CPU máxima medida', '2,84% sob carga (insignificante)'],
        ['Disco total', '~3,8 GB iniciais, crescimento gradual'],
        ['Portas expostas', '8090 (aplicação web) — MySQL 3306 não é afetado'],
        ['Dados do banco', 'Pasta nativa do Windows (ex: D:\\PostgreSQL\\data) — acesso direto'],
        ['Backup', 'pg_dump nativo + script .bat agendado no Task Scheduler'],
        ['Restauração', 'pg_restore nativo — comando único'],
        ['Inicialização automática', 'Sim — PostgreSQL como serviço Windows + containers Docker com auto-restart'],
        ['Comunicação com MySQL/farmácia', 'Nenhuma — sistemas completamente independentes'],
        ['Risco para o sistema atual', 'Nenhum — isolamento total por design'],
    ]
)

# Rodapé
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Relatório gerado para avaliação técnica da instalação do Docker e PostgreSQL no ambiente de produção.')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save(r'C:\Users\zkozt\Enc\Relatorio_Docker_Servidor.docx')
print('OK - Relatorio_Docker_Servidor.docx gerado com sucesso!')
